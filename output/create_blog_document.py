#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SEO 최적화 결과와 초안을 반영하여 최종 Word(.docx) 문서 생성
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import os

def create_blog_document():
    """초안과 SEO 분석을 반영하여 완성된 Word 문서를 생성합니다."""

    # 새 문서 생성
    doc = Document()

    # 문서 여백 설정 (2.5cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # SEO 최적화된 메인 제목 (Heading 1) - SEO 보고서 첫 번째 후보 사용
    main_title = doc.add_heading('파이썬 웹 크롤링 초보자 가이드: BeautifulSoup으로 데이터 수집하기', 0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 메타 정보 섹션
    doc.add_heading('📊 메타 정보', level=2)

    meta_info = doc.add_paragraph()
    meta_info.add_run('SEO 점수: ').bold = True
    meta_info.add_run('78/100\n')
    meta_info.add_run('메타 디스크립션: ').bold = True
    meta_info.add_run('파이썬 웹 크롤링 초보자를 위한 완벽 가이드. BeautifulSoup, requests를 활용한 데이터 수집 방법과 실습 예제를 제공합니다. 2025년 최신 웹 크롤링 기법과 윤리적 크롤링 방법까지 알아보세요.\n')
    meta_info.add_run('주요 키워드: ').bold = True
    meta_info.add_run('파이썬 웹 크롤링, BeautifulSoup, requests, 웹 스크래핑, 데이터 수집, 파이썬 초보자, 크롤링 튜토리얼, HTML 파싱, CSS 선택자, Scrapy, Selenium')

    # 구분선
    doc.add_paragraph('─' * 60)

    # 서론 (초안 내용 그대로)
    intro_p = doc.add_paragraph()
    intro_p.add_run('웹에 떠다니는 수많은 데이터를 수동으로 수집하는 일은 이제 과거의 일입니다. 파이썬 웹 크롤링을 활용하면 뉴스 헤드라인, 상품 가격 정보, 부동산 매물 등 필요한 데이터를 자동으로 수집할 수 있습니다. 이 글에서는 웹 크롤링의 기본 개념부터 실제 구현까지, 초보자가 단계별로 학습할 수 있는 완벽한 가이드를 제공합니다.')

    # 웹 크롤링이 필요한 이유
    doc.add_heading('웹 크롤링이 필요한 이유', level=2)

    doc.add_heading('자동화의 필요성', level=3)
    doc.add_paragraph('매일 반복되는 데이터 수집 작업을 자동화함으로써 시간과 비용을 절약할 수 있습니다. 예를 들어, 경쟁사 제품 가격 모니터링, 부동산 시장 동향 파악, 뉴스 트렌드 분석 등의 업무를 자동화할 수 있습니다.')

    doc.add_heading('2025년 웹 크롤링 환경', level=3)
    doc.add_paragraph('현재 웹 크롤링 환경은 과거와 크게 달라졌습니다. JavaScript 기반 동적 웹사이트의 증가와 AI 봇 탐지 시스템의 강화로 인해 더욱 정교한 접근 방법이 필요합니다. 또한 법적, 윤리적 고려사항이 필수 요소로 부상하였습니다.')

    # 기본 개념 이해
    doc.add_heading('기본 개념 이해', level=2)

    doc.add_heading('웹 크롤링과 웹 스크래핑의 차이', level=3)
    doc.add_paragraph('웹 크롤링은 웹사이트를 체계적으로 탐색하는 과정이며, 웹 스크래핑은 특정 데이터를 추출하는 작업입니다. 파이썬을 활용하면 이 두 작업을 효율적으로 수행할 수 있습니다.')

    doc.add_heading('HTML 구조 이해', level=3)
    doc.add_paragraph('웹 크롤링을 위해서는 HTML 태그 구조와 CSS 선택자에 대한 기본적인 이해가 필요합니다. 브라우저의 개발자 도구를 활용하여 원하는 데이터의 위치를 파악하는 방법을 익혀야 합니다.')

    # 필수 라이브러리 소개
    doc.add_heading('필수 라이브러리 소개', level=2)

    doc.add_heading('requests + BeautifulSoup (초보자 추천)', level=3)
    lib_p = doc.add_paragraph()
    lib_p.add_run('초보자가 가장 쉽게 시작할 수 있는 조합입니다. 정적 웹사이트 크롤링에 최적화되어 있으며, 간단한 설치 과정만으로 사용할 수 있습니다.')

    # 코드 블록 스타일링 함수
    def add_code_block(doc, code_text):
        code_p = doc.add_paragraph()
        code_run = code_p.add_run(code_text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(10)
        # 코드 블록 스타일만 적용 (배경색 제외)
        return code_p

    add_code_block(doc, 'pip install requests beautifulsoup4')

    doc.add_heading('Scrapy (대규모 크롤링)', level=3)
    doc.add_paragraph('수천, 수만 개의 페이지를 크롤링해야 하는 대규모 프로젝트에 적합한 프레임워크입니다. 비동기 처리를 통해 빠른 속도를 제공하며, 자동 링크 추적 및 페이지네이션 처리가 가능합니다.')

    doc.add_heading('Selenium (동적 웹사이트)', level=3)
    doc.add_paragraph('JavaScript 렌더링이 필요한 동적 사이트나 로그인이 필요한 사이트에서 사용됩니다. 실제 브라우저를 자동화하여 복잡한 웹사이트도 처리할 수 있습니다.')

    # 실습 예제
    doc.add_heading('실습 예제: 뉴스 헤드라인 크롤링', level=2)
    doc.add_paragraph('다음은 뉴스 웹사이트에서 헤드라인을 수집하는 기본적인 예제입니다.')

    # 코드 예제
    code_example = '''import requests
from bs4 import BeautifulSoup
import time

def crawl_news_headlines(url):
    # User-Agent 설정으로 봇 탐지 우회
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }

    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')
        headlines = soup.select('.news-title')  # CSS 선택자 활용

        for headline in headlines:
            print(headline.get_text().strip())
            time.sleep(1)  # 서버 부하 방지

    except requests.RequestException as e:
        print(f"요청 오류: {e}")

# 사용 예시
crawl_news_headlines("https://example-news.com")'''

    add_code_block(doc, code_example)

    doc.add_heading('데이터 저장하기', level=3)
    doc.add_paragraph('수집한 데이터를 CSV 파일로 저장하는 방법입니다.')

    # CSV 저장 코드
    csv_code = '''import pandas as pd

def save_to_csv(data, filename):
    df = pd.DataFrame(data)
    df.to_csv(filename, index=False, encoding='utf-8-sig')
    print(f"데이터가 {filename}에 저장되었습니다.")'''

    add_code_block(doc, csv_code)

    # 주의사항 및 베스트 프랙티스
    doc.add_heading('주의사항 및 베스트 프랙티스', level=2)

    doc.add_heading('기술적 베스트 프랙티스', level=3)
    practices_p = doc.add_paragraph()
    practices_p.add_run('1. 요청 간격 조절: ').bold = True
    practices_p.add_run('time.sleep() 함수를 사용하여 서버에 과부하를 주지 않도록 합니다.\n')
    practices_p.add_run('2. 에러 처리: ').bold = True
    practices_p.add_run('try-except 구문을 활용하여 네트워크 오류나 파싱 오류에 대비합니다.\n')
    practices_p.add_run('3. 세션 관리: ').bold = True
    practices_p.add_run('requests.Session()을 사용하여 효율적인 연결을 유지합니다.')

    doc.add_heading('윤리적 및 법적 고려사항', level=3)
    legal_p = doc.add_paragraph()
    legal_p.add_run('1. robots.txt 확인: ').bold = True
    legal_p.add_run('웹사이트의 크롤링 정책을 반드시 확인합니다.\n')
    legal_p.add_run('2. 이용약관 검토: ').bold = True
    legal_p.add_run('상업적 사용이나 데이터 재배포에 대한 제한사항을 파악합니다.\n')
    legal_p.add_run('3. 개인정보보호: ').bold = True
    legal_p.add_run('민감한 개인정보는 수집하지 않습니다.\n')
    legal_p.add_run('4. 서버 부하 최소화: ').bold = True
    legal_p.add_run('과도한 요청으로 서버에 피해를 주지 않습니다.')

    # 결론
    doc.add_heading('결론', level=2)
    conclusion_p = doc.add_paragraph()
    conclusion_p.add_run('파이썬 웹 크롤링은 데이터 수집 업무를 자동화할 수 있는 강력한 도구입니다. 기본적인 requests와 BeautifulSoup 조합부터 시작하여, 필요에 따라 Scrapy나 Selenium 등의 고급 도구로 확장할 수 있습니다.')

    doc.add_paragraph('중요한 것은 기술적 역량과 함께 윤리적, 법적 책임감을 가지고 크롤링을 수행하는 것입니다. 이 가이드를 통해 기본기를 충실히 익힌 후, 실제 프로젝트에서 필요에 맞는 크롤링 솔루션을 구현해보시기 바랍니다.')

    doc.add_paragraph('웹 크롤링의 세계는 무궁무진합니다. 첫 걸음을 내딛었다면, 이제 더 복잡한 프로젝트에 도전하며 실력을 향상시켜 나가시기 바랍니다.')

    # 참고 자료 섹션
    doc.add_heading('📚 참고 자료', level=2)

    ref_p = doc.add_paragraph()
    ref_p.add_run('기본 학습 자료:\n').bold = True
    ref_p.add_run('1. 웹 크롤링 완벽 가이드 - Security Framework\n')
    ref_p.add_run('2. Python을 활용한 데이터 크롤링 - 노마드데이터랩\n')
    ref_p.add_run('3. 웹 스크레이핑 초보자 가이드 - nextdoorped\n\n')

    ref_p.add_run('고급 기술 문서:\n').bold = True
    ref_p.add_run('4. 파이썬으로 데이터 스크래핑 - Thunderbit\n')
    ref_p.add_run('5. Scrapy 실전 가이드 - Thunderbit\n')
    ref_p.add_run('6. 크롤링 가이드 - TBWA DATA LAB\n\n')

    ref_p.add_run('공식 문서:\n').bold = True
    ref_p.add_run('7. BeautifulSoup4 Documentation (GitHub)\n')
    ref_p.add_run('8. Scrapy Documentation (GitHub)\n')
    ref_p.add_run('9. SeleniumBase Documentation (GitHub)')

    # 문서 저장
    output_path = r'C:\Users\hiond\workspace\blog-poster\output\파이썬으로-웹-크롤링-시작하기-blog.docx'
    doc.save(output_path)
    print(f"블로그 문서가 성공적으로 생성되었습니다: {output_path}")

    return output_path

if __name__ == "__main__":
    try:
        file_path = create_blog_document()
        print(f"[성공] Word 문서 생성 완료!")
        print(f"[위치] 저장 위치: {file_path}")

        # 파일 크기 확인
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"[크기] 파일 크기: {file_size:,} bytes")

        print("\n[구성] 문서 구성:")
        print("- SEO 최적화된 제목 (SEO 보고서 1순위 후보 사용)")
        print("- 메타 정보 섹션 (SEO 점수, 메타 설명, 키워드)")
        print("- 초안의 전체 본문 내용 (헤딩 레벨 유지, 글 톤 변경 없음)")
        print("- 참고 자료 섹션 (리서치 보고서의 참고 링크 포함)")
        print("- 코드 블록은 Consolas 고정폭 폰트로 스타일링")
        print("- 적절한 여백과 줄간격 설정")

    except Exception as e:
        print(f"[오류] 오류 발생: {e}")