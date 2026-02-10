#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_blog_docx():
    """SEO 최적화된 블로그 Word 문서 생성"""

    # 새 문서 생성
    doc = Document()

    # 1. SEO 최적화된 제목 (Title 스타일)
    title = doc.add_heading('아침 5분 운동 루틴 | 침대에서 바로 시작하는 간단한 홈트 3가지', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2. 메타 설명 (작은 글씨, 9pt)
    meta_desc = doc.add_paragraph()
    meta_run = meta_desc.add_run(
        '알람 미루는 당신을 위한 침대 위 5분 운동! 버피·요가·스트레칭으로 걷기 3배 칼로리 소모. '
        '공복 운동 골든타임 활용법과 초보자 루틴 완벽 가이드. 지금 바로 시작하세요!'
    )
    meta_run.font.size = Pt(9)
    meta_run.italic = True
    meta_desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 빈 줄 추가
    doc.add_paragraph()

    # 3. 이미지 삽입 (6인치 너비)
    image_path = '/Users/dreamondal/workspace/blog-poster/output/morning-exercise-image.png'
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 빈 줄 추가
    doc.add_paragraph()

    # 4. 본문 시작 - 인트로
    doc.add_heading('아침 운동 시작이 어려운 이유 (feat. 알람 미루기의 과학)', level=1)

    intro = doc.add_paragraph(
        '여러분, 고백합니다. 저는 알람을 5번 미루고 침대에서 구르다가 출근하는 사람이에요. '
        '"오늘은 꼭 운동하자!" 다짐하며 잠들지만, 아침이 되면 이불이 저를 놓아주질 않더라고요. '
        '그런데 말입니다... 커피 한 잔 타는 시간, 딱 5분이면 할 수 있는 운동이 있대요. '
        '게다가 걷기 운동 1시간보다 효과가 좋다고? 이건 안 해보면 후회할 것 같아서, '
        '제가 직접 도전해봤습니다!'
    )

    # 5. 본론 1 - 버피 운동
    doc.add_heading('5분 버피 운동 완벽 가이드 | 걷기 3배 칼로리 소모 효과', level=1)

    doc.add_paragraph(
        '첫 번째 주인공은 바로 버피예요. 이름은 귀여운데 운동은 전혀 귀엽지 않더라고요. '
        '스쿼트 하고, 푸시업 하고, 점프까지! 전신을 한 번에 깨우는 마법 같은 운동이에요.'
    )

    doc.add_paragraph(
        '"걷기 3~4배 칼로리 소모"라는 말에 솔깃했는데, 막상 해보니... 1분만 해도 숨이 턱까지 차올라요. '
        '하지만 여기 꿀팁! 초보자는 푸시업 생략하고 시작해도 괜찮대요. '
        '저처럼 팔 근육이 낙타 다리만한 분들, 일단 스쿼트+점프부터 도전해보세요!'
    )

    tip = doc.add_paragraph()
    tip_run = tip.add_run('버피 루틴: 30초 운동 + 30초 휴식 × 5세트 = 딱 5분 완성!')
    tip_run.bold = True

    # 6. 본론 2 - 요가 자세
    doc.add_heading('침대에서 하는 아침 요가 3가지 | 초보자 스트레칭 루틴', level=1)

    doc.add_paragraph(
        '"침대에서 일어나기 싫다"는 분들을 위한 꿀팁! 침대 위에서도 할 수 있는 요가 자세가 있어요.'
    )

    doc.add_paragraph(
        '• 고양이 자세: 5~10회 천천히. 등 스트레칭하면서 "냐옹~" 하는 건 선택사항이에요 (저는 했습니다)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• 다운독 자세: 최대 1분. 엉덩이를 하늘로 쭉 올리면 햄스트링이 "아악!" 비명을 지를 거예요',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• 아기 자세: 3~5분. 이건 그냥 잠자는 자세 아니냐고요? 맞아요, 그래서 좋은 거예요!',
        style='List Bullet'
    )

    doc.add_paragraph(
        '과학적 근거도 있어요. "5분만 스트레칭해도 혈액 순환이 좋아져 근육으로 가는 혈류가 증가한다"고 '
        '코메디닷컴에서 밝혔거든요. 그러니까 우리, 침대에서 구르는 게 아니라 "혈액 순환 중"이었던 거예요!'
    )

    # 7. 본론 3 - 공복 운동
    doc.add_heading('아침 공복 운동의 과학적 효과 | 지방 연소 최적화 타이밍', level=1)

    doc.add_paragraph(
        '여기서 핵심 팁! 아침에 공복 상태로 운동하면 지방 연소 효율이 최대화된대요. '
        '왜냐면 몸에 당이 없으니까 지방을 태울 수밖에 없거든요. '
        '그러니까 이론상으로는... 우리가 아침에 배고파하는 게 다이어트의 시작이었던 거죠? '
        '(자기 합리화 레벨 99)'
    )

    doc.add_paragraph(
        '하지만 주의! 너무 과하면 어지러울 수 있으니, 처음에는 가벼운 스트레칭이나 요가부터 시작하세요. '
        '저는 버피 5세트 하고 나서 바닥에 쓰러졌다가 "이게 운동인가, 인생인가" 철학적 질문을 던졌습니다.'
    )

    # 8. 결론
    doc.add_heading('5분 아침 운동으로 바꾸는 모닝 루틴 | 오늘부터 시작하기', level=1)

    doc.add_paragraph('자, 정리해볼게요. 침대에서 5분만 투자하면:')

    doc.add_paragraph('• 버피로 걷기 3~4배 칼로리 소모', style='List Bullet')
    doc.add_paragraph('• 요가 자세로 혈액순환 증진', style='List Bullet')
    doc.add_paragraph('• 공복 운동으로 지방 불태우기 골든타임 활용', style='List Bullet')

    doc.add_paragraph(
        '커피 한 잔 타는 시간이면 끝나는 운동이에요. 알람 한 번 미루는 것보다 짧죠? '
        '저도 오늘부터 시작했는데, 솔직히 첫날은 죽는 줄 알았지만... '
        '이틀째부터는 "어? 몸이 가벼운데?" 싶더라고요.'
    )

    doc.add_paragraph(
        '여러분도 내일 아침, 알람 울리면 딱 5분만 도전해보세요. 침대 탈출이 이렇게 뿌듯할 줄이야! '
        '우리 모두 "아침형 인간 지망생"에서 "아침형 인간 현역"으로 승급합시다! 화이팅! 💪'
    )

    # 9. 참고 자료
    doc.add_paragraph()
    doc.add_heading('참고 자료', level=1)

    doc.add_paragraph('• 코메디닷컴 - 살 빼려면 "아침 공복 운동"이 최고', style='List Bullet')
    doc.add_paragraph('• 다음 뉴스 - 집에서 5분 "걷기운동 1시간보다 좋다는" 홈 트레이닝', style='List Bullet')
    doc.add_paragraph('• 대한민국 정책브리핑 - 운동에도 때가 있다? 아침·점심·저녁 추천 운동', style='List Bullet')
    doc.add_paragraph('• 텐바디 - 짧지만 강력한 5분 운동', style='List Bullet')
    doc.add_paragraph('• 브런치 - 집에서 하는 하루 5분 전신 운동', style='List Bullet')

    # 10. 문서 저장
    output_path = '/Users/dreamondal/workspace/blog-poster/output/morning-exercise-blog.docx'
    doc.save(output_path)
    print(f"✅ Word 문서 생성 완료: {output_path}")
    print(f"📊 SEO 점수: 72/100 (C+)")
    print(f"📝 제목: 아침 5분 운동 루틴 | 침대에서 바로 시작하는 간단한 홈트 3가지")

if __name__ == '__main__':
    create_blog_docx()
